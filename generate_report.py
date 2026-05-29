from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_chapter_heading(doc, text):
    heading = doc.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def create_report():
    doc = Document()
    
    # Front Page
    for _ in range(5): doc.add_paragraph()
    title = doc.add_heading('PROJECT REPORT\n\nPARENT-TEACHER COMMUNICATION PORTAL:\nDOCKERIZATION & CI/CD PIPELINE IMPLEMENTATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(15): doc.add_paragraph()
    doc.add_paragraph('Submitted by: Abhay Rawat').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_text = (
        "Chapter 1 — Introduction\n"
        "   1.1 Project Overview\n"
        "   1.2 Problem Statement\n"
        "   1.3 Objectives\n"
        "   1.4 Scope\n\n"
        "Chapter 2 — System Requirements\n"
        "   2.1 Hardware Requirements\n"
        "   2.2 Software Requirements\n"
        "   2.3 Technology Stack Justification\n\n"
        "Chapter 3 — Workflow\n"
        "   3.1 Architecture Overview\n"
        "   3.2 Step-by-Step Process\n\n"
        "Chapter 4 — Implementation\n"
        "   4.1 Setup Steps\n"
        "   4.2 Configurations\n"
        "   4.3 Commands\n"
        "   4.4 Dockerfile (Backend & Frontend)\n"
        "   4.5 CI/CD Pipeline (GitHub Actions / Jenkinsfile)\n\n"
        "Chapter 5 — Results and Output\n"
        "   5.1 Successful Pipeline Build\n"
        "   5.2 Docker Image Creation\n"
        "   5.3 Running Containers\n"
        "   5.4 Deployed Application\n\n"
        "Chapter 6 — Future Scope\n"
        "Chapter 7 — Conclusion"
    )
    doc.add_paragraph(toc_text)
    doc.add_page_break()
    
    # Chapter 1
    add_chapter_heading(doc, 'Chapter 1 — Introduction')
    doc.add_heading('1.1 Project Overview', level=2)
    doc.add_paragraph("The Parent-Teacher Communication Portal is a comprehensive, full-stack web application designed to bridge the communication gap between educational institutions and parents. Built with a Laravel (PHP) backend and a React (Vite) frontend, the platform offers real-time notifications, grade tracking, assignment submissions, and a multi-student dashboard for parents. As modern software engineering moves rapidly towards cloud-native architectures, this phase of the project focuses intensely on DevOps practices. Specifically, it involves the complete dockerization of the application stack and the establishment of a robust Continuous Integration and Continuous Deployment (CI/CD) pipeline.")
    doc.add_paragraph("By integrating containerization, the project ensures that the application runs uniformly across all environments—from the developer's local machine to testing servers and production environments. The CI/CD pipeline automates testing and deployment processes, drastically reducing the manual overhead required to release new features and ensuring high code quality.")
    
    doc.add_heading('1.2 Problem Statement', level=2)
    doc.add_paragraph("Historically, deploying full-stack applications manually has been fraught with challenges. Developers frequently encounter the 'it works on my machine' syndrome, where software behaves differently in production than it does in development due to mismatched environment variables, missing dependencies, or differing operating system configurations.")
    doc.add_paragraph("Furthermore, the manual testing, linting, and building of the frontend and backend repositories are incredibly time-consuming and error-prone. Without an automated pipeline, releasing updates requires developers to manually execute test suites, compile static assets, manage database migrations, and configure web servers. This lack of automation bottlenecks the development lifecycle and increases the risk of deploying broken code to production.")
    
    doc.add_heading('1.3 Objectives', level=2)
    doc.add_paragraph("The primary objectives of this project phase include:")
    doc.add_paragraph("1. Containerization: Isolate the Laravel backend and React frontend into lightweight, executable Docker containers.")
    doc.add_paragraph("2. Orchestration: Utilize Docker Compose to manage multi-container deployments, seamlessly linking the frontend, backend, web servers, and MySQL database.")
    doc.add_paragraph("3. Automation: Establish a CI/CD pipeline that automatically triggers upon code commits to test, build, and package the application.")
    doc.add_paragraph("4. Consistency: Eliminate environment discrepancies by encapsulating all system-level dependencies within the Dockerfiles.")
    
    doc.add_heading('1.4 Scope', level=2)
    doc.add_paragraph("The scope of this project encompasses writing detailed, multi-stage Dockerfiles for both the PHP backend and Node/React frontend. It includes configuring Nginx web servers to properly route API requests and serve Single Page Application (SPA) assets. Additionally, the scope involves writing a root docker-compose.yml file for local orchestration and authoring a comprehensive CI/CD pipeline script (using GitHub Actions, acting equivalently to a Jenkinsfile) that automatically executes PHPUnit tests, ESLint checks, and Docker builds.")
    doc.add_page_break()
    
    # Chapter 2
    add_chapter_heading(doc, 'Chapter 2 — System Requirements')
    doc.add_heading('2.1 Hardware Requirements', level=2)
    doc.add_paragraph("To effectively run the containerized application and the Docker daemon, the following hardware specifications are required:")
    doc.add_paragraph("- Processor: Intel Core i5 / AMD Ryzen 5 or equivalent (Minimum 2.5 GHz multi-core processor for concurrent container execution).")
    doc.add_paragraph("- RAM: Minimum 8 GB (16 GB is highly recommended for running multiple Docker containers alongside an IDE and browser).")
    doc.add_paragraph("- Storage: 20 GB of free disk space (Solid State Drive (SSD) is preferred to ensure fast image pulling and container startup).")
    doc.add_paragraph("- Network: Broadband internet connection for pulling base Docker images and pushing to repositories.")
    
    doc.add_heading('2.2 Software Requirements', level=2)
    doc.add_paragraph("- Operating System: Linux (Ubuntu 20.04+, Debian), macOS (10.15+), or Windows 10/11 (with Windows Subsystem for Linux 2 - WSL2 enabled).")
    doc.add_paragraph("- Container Engine: Docker Engine (v24.0+) and Docker Compose plugin.")
    doc.add_paragraph("- Version Control: Git.")
    doc.add_paragraph("- Source Code Editor: Visual Studio Code or PhpStorm.")
    
    doc.add_heading('2.3 Technology Stack Justification', level=2)
    doc.add_paragraph("Docker was chosen as the containerization platform due to its industry-standard status, immense community support, and lightweight resource footprint compared to traditional virtual machines. Nginx was selected over Apache due to its asynchronous, event-driven architecture, which makes it exceptionally fast at serving static React assets and proxying requests to PHP-FPM. GitHub Actions was utilized for the CI/CD pipeline because it natively integrates with the version control system, eliminating the need to host and maintain a separate Jenkins master node, while providing identical pipeline capabilities.")
    doc.add_page_break()
    
    # Chapter 3
    add_chapter_heading(doc, 'Chapter 3 — Workflow')
    doc.add_heading('3.1 Architecture Overview', level=2)
    doc.add_paragraph("The system architecture relies on a microservices-inspired setup orchestrated by Docker Compose. The architecture is divided into four distinct services interacting over an isolated bridge network: 'parent-teacher-app_app-network'.")
    doc.add_paragraph("1. Database Service (db): Runs MySQL 8.0, storing persistent application data mounted to a local Docker volume.")
    doc.add_paragraph("2. Backend Service (backend): Runs PHP 8.4-FPM, executing Laravel business logic and querying the database.")
    doc.add_paragraph("3. Web Server Service (web): Runs Nginx, acting as a reverse proxy that receives HTTP requests on port 8000 and forwards PHP scripts to the backend service via FastCGI.")
    doc.add_paragraph("4. Frontend Service (frontend): Runs Nginx, serving the compiled static assets of the React application on port 3000.")
    
    doc.add_heading('3.2 Step-by-Step Process', level=2)
    doc.add_paragraph("The deployment and continuous integration workflow proceeds systematically through the following steps:")
    doc.add_paragraph("Step 1: Containerize Backend. A Dockerfile is crafted for the Laravel application using the official php:8.4-fpm image. System-level dependencies required by Laravel (such as libzip, libxml, and curl) are installed. PHP extensions (pdo_mysql, mbstring, gd) are compiled and enabled. Finally, Composer is utilized to install PHP dependencies in a non-interactive mode without executing scripts prematurely.")
    doc.add_paragraph("Step 2: Containerize Frontend. A multi-stage Dockerfile is constructed for React. In the 'build-stage', a Node.js image installs NPM dependencies and compiles the Vite application into static HTML/CSS/JS files. In the 'production-stage', an Alpine Linux Nginx image simply copies these compiled files and serves them. This multi-stage approach ensures the final image is extremely small and secure, containing no Node.js runtime or source code.")
    doc.add_paragraph("Step 3: Setup Web Server Configuration. Custom Nginx configuration files (default.conf) are written for both the frontend and backend. The backend configuration uses the 'fastcgi_pass' directive to route traffic to port 9000 of the PHP container. The frontend configuration uses 'try_files' to ensure React Router handles client-side routing properly by falling back to index.html.")
    doc.add_paragraph("Step 4: Orchestration Setup. The docker-compose.yml file is written to define all four services, map host ports to container ports (3000 for frontend, 8000 for backend, 3306 for DB), create isolated volumes for database persistence, and inject environment variables (like DB_PASSWORD).")
    doc.add_paragraph("Step 5: CI/CD Setup. A pipeline file (.github/workflows/ci.yml) is authored. It is triggered on 'push' and 'pull_request' events to the main branch. The pipeline checks out the code, provisions a testing environment (including setting up PHP 8.4 and Node 22), creates an in-memory SQLite database, runs the Laravel PHPUnit test suite, executes the frontend ESLint checks, and ultimately tests building both Docker images to ensure no compilation errors exist.")
    doc.add_paragraph("Step 6: Local Execution & Deployment. The developer executes 'docker compose up -d --build'. Docker pulls necessary base images, executes the Dockerfiles, links the containers via a bridge network, and starts the services. Post-startup, database migrations are triggered manually via 'docker compose exec'.")
    doc.add_page_break()
    
    # Chapter 4
    add_chapter_heading(doc, 'Chapter 4 — Implementation')
    doc.add_heading('4.1 Setup Steps', level=2)
    doc.add_paragraph("To set up the project locally, the repository must first be cloned. Ensure Docker Desktop is running. No local installations of PHP, Composer, or Node.js are strictly necessary, as all execution happens within the containers. The environment variables are passed automatically via Docker Compose.")
    
    doc.add_heading('4.2 Configurations', level=2)
    doc.add_paragraph("The backend Nginx configuration (backend/nginx/default.conf) intercepts requests to port 80 and processes PHP files:")
    doc.add_paragraph("location ~ \.php$ {\n    fastcgi_pass backend:9000;\n    fastcgi_index index.php;\n    include fastcgi_params;\n    fastcgi_param SCRIPT_FILENAME $document_root$fastcgi_script_name;\n}")
    
    doc.add_heading('4.3 Commands', level=2)
    doc.add_paragraph("The following terminal commands are integral to managing the Dockerized application:")
    doc.add_paragraph("- docker compose up -d --build : Builds fresh images and starts all services in detached mode.")
    doc.add_paragraph("- docker compose exec backend php artisan migrate : Connects to the running backend container and executes database migrations.")
    doc.add_paragraph("- docker compose logs -f : Streams real-time logs from all containers.")
    doc.add_paragraph("- docker compose down : Stops and removes all containers and networks.")
    
    doc.add_heading('4.4 Dockerfile Definitions', level=2)
    doc.add_paragraph("The Backend Dockerfile heavily customizes the PHP-FPM image:")
    doc.add_paragraph("FROM php:8.4-fpm\nRUN apt-get update && apt-get install -y git curl libpng-dev zip unzip\nRUN docker-php-ext-install pdo_mysql mbstring gd\nCOPY --from=composer:latest /usr/bin/composer /usr/bin/composer\nWORKDIR /var/www\nCOPY . /var/www\nRUN composer install --no-interaction --prefer-dist --optimize-autoloader --no-scripts\nRUN chown -R www-data:www-data /var/www\nEXPOSE 9000\nCMD [\"php-fpm\"]")
    
    doc.add_paragraph("\nThe Frontend Dockerfile employs a multi-stage build:")
    doc.add_paragraph("FROM node:22-alpine as build-stage\nWORKDIR /app\nCOPY package*.json ./\nRUN npm ci\nCOPY . .\nRUN npm run build\n\nFROM nginx:alpine as production-stage\nCOPY nginx/default.conf /etc/nginx/conf.d/default.conf\nCOPY --from=build-stage /app/dist /usr/share/nginx/html\nEXPOSE 80\nCMD [\"nginx\", \"-g\", \"daemon off;\"]")
    
    doc.add_heading('4.5 CI/CD Pipeline (Jenkinsfile Equivalent)', level=2)
    doc.add_paragraph("The Continuous Integration pipeline automates quality assurance. Below is a conceptual Jenkinsfile mapping exactly to the GitHub Actions implementation utilized in this project:")
    doc.add_paragraph("pipeline {\n    agent any\n    stages {\n        stage('Checkout') {\n            steps { checkout scm }\n        }\n        stage('Backend Tests') {\n            steps {\n                sh 'composer install'\n                sh 'touch database/database.sqlite'\n                sh 'php artisan migrate --force'\n                sh 'php artisan test'\n            }\n        }\n        stage('Frontend Lint & Build') {\n            steps {\n                sh 'npm ci'\n                sh 'npm run lint'\n                sh 'npm run build'\n            }\n        }\n        stage('Build Docker Images') {\n            steps {\n                sh 'docker build -t app-backend ./backend'\n                sh 'docker build -t app-frontend ./frontend'\n            }\n        }\n    }\n}")
    doc.add_page_break()
    
    # Chapter 5
    add_chapter_heading(doc, 'Chapter 5 — Results and Output')
    doc.add_paragraph("The implementation successfully transformed the Parent-Teacher Portal into a robust, cloud-ready application. Below are the documented outputs (Please insert corresponding screenshots in the final document):")
    
    doc.add_heading('5.1 Successful Pipeline Build', level=2)
    doc.add_paragraph("Upon pushing code to the repository, the CI pipeline triggers automatically. The terminal logs output success messages for the PHPUnit tests (confirming models and endpoints function correctly against the SQLite testing database) and the ESLint checks (confirming frontend code quality).")
    doc.add_paragraph("\n[ Insert Screenshot: CI/CD Pipeline passing with green checkmarks ]\n")
    for _ in range(5): doc.add_paragraph()
    
    doc.add_heading('5.2 Docker Image Creation', level=2)
    doc.add_paragraph("The Docker daemon successfully executes both Dockerfiles. The multi-stage frontend build completes by producing an image sized under 30MB, significantly optimizing deployment times.")
    doc.add_paragraph("\n[ Insert Screenshot: Terminal output showing 'docker build' completing layers successfully ]\n")
    for _ in range(5): doc.add_paragraph()
    
    doc.add_heading('5.3 Running Containers', level=2)
    doc.add_paragraph("Executing 'docker compose ps' reveals four active containers. The database container binds to an internal network, safely isolating it from external traffic, while the Nginx containers expose ports 3000 and 8000 to the host.")
    doc.add_paragraph("\n[ Insert Screenshot: Terminal showing running containers and mapped ports ]\n")
    for _ in range(5): doc.add_paragraph()
    
    doc.add_heading('5.4 Deployed Application', level=2)
    doc.add_paragraph("The frontend application renders seamlessly on localhost:3000. It successfully makes cross-origin requests to the backend API on localhost:8000, establishing a fully functional full-stack loop.")
    doc.add_paragraph("\n[ Insert Screenshot: Web browser showing the Parent-Teacher Dashboard loaded successfully ]\n")
    doc.add_page_break()
    
    # Chapter 6
    add_chapter_heading(doc, 'Chapter 6 — Future Scope')
    doc.add_paragraph("While this phase successfully containerized the application and established continuous integration, several advanced DevOps practices can be integrated in the future to further enhance the system:")
    doc.add_paragraph("1. Container Registry Integration: The CI/CD pipeline can be expanded to automatically push successful Docker image builds to a remote registry such as Docker Hub, GitHub Container Registry (GHCR), or Amazon Elastic Container Registry (ECR).")
    doc.add_paragraph("2. Kubernetes Orchestration: For enterprise-level scalability, the docker-compose setup can be translated into Kubernetes manifests (Deployments, Services, Ingress). This would allow the application to auto-scale horizontally across multiple nodes based on traffic spikes during school hours.")
    doc.add_paragraph("3. Automated End-to-End (E2E) Testing: Tools like Cypress or Selenium can be added to the CI pipeline. Unlike unit tests, E2E tests will spin up the entire Docker orchestration temporarily and simulate real user clicks to ensure the UI behaves exactly as expected.")
    doc.add_paragraph("4. Security Scanning: Integration of automated vulnerability scanners (like Trivy or SonarQube) in the CI pipeline to scan the Docker images and source code for known security vulnerabilities before deployment.")
    doc.add_paragraph("5. Zero-Downtime Deployments: Implementing blue-green deployments or rolling updates to ensure that parents and teachers experience zero downtime when the portal is updated to a new version.")
    doc.add_page_break()
    
    # Chapter 7
    add_chapter_heading(doc, 'Chapter 7 — Conclusion')
    doc.add_paragraph("In conclusion, the project successfully achieved its primary objective of modernizing the deployment and testing architecture of the Parent-Teacher Communication Portal. By transitioning from a traditional, manual deployment strategy to a containerized Docker architecture, we have inherently resolved the persistent issue of environment inconsistencies. Developers can now onboard in minutes simply by running 'docker compose up', drastically accelerating development velocity.")
    
    doc.add_paragraph("Furthermore, the implementation of the Continuous Integration pipeline guarantees software reliability. The automated execution of PHPUnit tests and ESLint prevents degraded code from ever merging into the main branch. The inclusion of Docker image building within the pipeline acts as an ultimate verification step, ensuring that the application remains buildable at all times.")
    
    doc.add_paragraph("Through this project, profound technical proficiencies were developed. Core technologies mastered include Docker containerization, multi-stage image optimization, Docker Compose network orchestration, Nginx reverse proxy configuration, and CI/CD pipeline automation utilizing GitHub Actions and Jenkins principles.")
    
    doc.add_paragraph("Ultimately, the final outcome is a highly resilient, scalable, and automated full-stack application environment. This solid foundation ensures that as the Parent-Teacher Portal scales to support thousands of students and teachers, the infrastructure will reliably support rapid iterations and deployments without sacrificing stability.")
    
    doc.save('Project_Report.docx')

if __name__ == '__main__':
    create_report()
